#!/usr/bin/env python
# coding: utf-8

# ### Creates figure names each genus and recreate the genus table with the appropriate figure link.

import docx
import pandas as pd


def get_para_data(output_doc_name, paragraph):
    """
    For copying a paragraph with all its characteristics
    No true copy command exist, so manually it is!
    All runs of a paragraph need to be copied individually
    
    Input:
        paragraph: The paragraph to transfer
        output_doc_name : The Document in which the paragraph is to be copied
    """
    # Creating a new paragraph (output) in the targeted Document
    output_para = output_doc_name.add_paragraph()
    # Copying all the runs of the input paragraph to the output paragraph
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        output_run.bold = run.bold
        output_run.italic = run.italic
        output_run.underline = run.underline
        output_run.font.color.rgb = run.font.color.rgb
        output_run.style.name = run.style.name
        output_run.font.size = run.font.size
        output_run.font.superscript = run.font.superscript
    # Copying Paragraph's format data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
    output_para.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
    #output_para.paragraph_format.element = paragraph.paragraph_format.element
    output_para.paragraph_format.keep_together = paragraph.paragraph_format.keep_together
    output_para.paragraph_format.keep_with_next = paragraph.paragraph_format.keep_with_next
    output_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
    output_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
    output_para.paragraph_format.line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
    output_para.paragraph_format.page_break_before = paragraph.paragraph_format.page_break_before
    output_para.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
    output_para.paragraph_format.space_after = paragraph.paragraph_format.space_after
    output_para.paragraph_format.space_before = paragraph.paragraph_format.space_before
    #output_para.paragraph_format.tab_stops = paragraph.paragraph_format.tab_stops
    output_para.paragraph_format.widow_control = paragraph.paragraph_format.widow_control


def copy_doc(to_doc, from_doc):
    """
    Copy a whole document (copy each paragraph)
    
    Inputs:
        to_doc: Document in which all paragraph are to be copied
                New paragraph are added at the end of the already existing ones
        from_doc: Document where all paragraph are to be copied.
    """
    for para in from_doc.paragraphs:
        get_para_data(to_doc, para)


### Script starts here ------------------------------------------------------

if __name__ == "__main__":
    
    # open csv with the genus list and the number of regular trees and short trees
    df = pd.read_csv('tree_count.csv', index_col='genus')
    
    # Initiating new column with the figure number
    df['figure'] = 0
    
    # Create the final Document from the template
    final_doc = docx.Document('template.docx')
    
    # Figure counter
    figure = 1
    fig_suffix = ['a','b','c','d','e','f']
    
    for genus in df.index:
        # In case of a single tree
        if df['trees'][genus] == 1:
            # Create a temp doc containing one figure, created with the template
            temp_doc = docx.Document('template.docx')
            # Modify
            temp_doc.paragraphs[0].runs[0].text = f'Figure {figure}.'
            temp_doc.paragraphs[0].runs[4].text = genus
            df['figure'][genus] = str(figure)
            # Save
            copy_doc(final_doc, temp_doc)
        # In case of more than one tree; add letter to the figure number
        else:
            fig_list = []
            suffix_index = 0
            for _ in range(df['long'][genus]):
                # Create a temp doc containing one figure, created with the template
                temp_doc = docx.Document('template.docx')
                # Modify
                temp_doc.paragraphs[0].runs[0].text = f'Figure {figure}{fig_suffix[suffix_index]}.'
                temp_doc.paragraphs[0].runs[4].text = genus
                fig_list.append(f'{figure}{fig_suffix[suffix_index]}')
                suffix_index += 1
                # Save
                copy_doc(final_doc, temp_doc)
            for _ in range(df['short'][genus]):
                # Create a temp doc containing one figure, created with the template
                temp_doc = docx.Document('template.docx')
                # Modify
                temp_doc.paragraphs[0].runs[0].text = f'Figure {figure}{fig_suffix[suffix_index]}.'
                temp_doc.paragraphs[0].runs[4].text = genus
                temp_doc.paragraphs[0].runs[5].text = temp_doc.paragraphs[0].runs[5].text + ' Individual analysis of short sequences.'
                fig_list.append(f'{figure}{fig_suffix[suffix_index]}')
                suffix_index += 1
                # Save
                copy_doc(final_doc, temp_doc)
            # Save figures as a list    
            df['figure'].loc[genus] = fig_list
        figure += 1
    
    
    # Save the final doc
    final_doc.save('output.docx')
    
    
    def changefigure(fig):
        # Change a figure list into a readable string
        if type(fig) is list:
            return ', '.join(fig)
        else:
            return fig
    
    # Change figure lists into a readable string
    df['figure'] = df['figure'].apply(changefigure)
    
    # Save genus list with assosiated figure number
    df.to_csv('tree_count_figure.csv')
    
    # Load the sample table
    dfsamples = pd.read_csv('TotalTableCSV_MOD.csv')
    
    # Create a new figure column, initiate with empty string.
    dfsamples['Figure'] = ''
    
    # Copy the appropriate figure number to each samples
    for index in dfsamples.index:
        dfsamples.loc[index]['Figure'] = df.loc[dfsamples.loc[index]['Genus']]['figure']
    
    # Save the table with the added figure column
    dfsamples.to_csv('TotalTableCSV_MOD_MOD.csv')






