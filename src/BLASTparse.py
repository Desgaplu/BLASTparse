#! python3
# -*- coding: utf-8

"""
BLASTparse.py
Created on Mon Dec 14 14:54:15 2020

Descrition:
Parse a Single-file JSON BLAST output for the best match of each sequence and
uses the fasta sequences of the samples to classify them in folder names with
the closest genus identified with the BLAST search.

Output a CSV file containing the sample name, samples seq length,
closest species, %identity, and other relevant informations.
It then create a folder for each genus and align each sample to the closest
TYPE strains and build a phylogenetic tree.

@author: pldesgagne
"""

import tkinter as tk  # using the open filedialog
from tkinter import filedialog  # using the open filedialog
import ntpath  # get the filename from path string
import json
import os
import pandas as pd
from sys import exit
from Bio.Align.Applications import MuscleCommandline
import docx

# Muscle.exe need to be in the same folder as this script
# from http://www.drive5.com/muscle/downloads.htm

# MEGACC (MEGA command console) need to be installed on computer
# The script use megacc command line from os.system() to create NJ trees
# The parameters of the tree creation are saved in file infer_NJ_nucleotide.mao
# This file is created with MEGAX in prototype mode


"""
TODO:

    -Ktinker GUI? Meh...
    
    - When merging the TYPE strain sequences between samples of same genus,
        if a duplicate is encountered, select the longest sequences.
        This way, we dont need to fetch the type sequences again because
        they might be to short for other external analysis.

"""

# Initialise tkinter to enable the uses of filedialog
root = tk.Tk()
root.withdraw()  # Prevent a empty window to be opened


class BlastParse():
    """
    Parse the multiple-json output of BLASTn with multiple 16S or 26S
    rDNA sequences and arrange them in a table.

    It is also possible to align all the sequences of the same genus
    with their best hits and create a NJ tree with the output. Those are
    saved in individual folder (named with the genus).

    Parameters
        ----------
        json_list : DICT
            Contain a list of the name of each json file.
        jsonfolder : STR
            Path to the json files.
        fasta_seqs : DICT
            Fasta sequences of the BLASTed samples.
            Key= Sample name STR
            Value= Fasta seq STR
        species_name_filter : LIST of STR
            keyword for the clean up of species names
        min_hits_filter : STR
            Minimum number of possible hits
        max_hits_filter : STR
            Maximum number of possible hits
        min_percentage : STR
            Minimum id percentage accepted for a hit
        best_hits_span : FLOAT
            Id percentage distance away from best hit for acceptable top hits
        short_seq_len : INT
            Len of the sequence under which the alignement is done separately.
        results : List of list
            Compiled results from all JSON files.
    """


    def __init__(self, json_list, jsonfolder, fasta_data):
        """
        Create the parser with default parameters.

        Parameters
        ----------
        json_list : DICT
            Contain a list of the name of each json file.
        jsonfolder : STR
            Path to the json files.
        fasta_seqs : STR
            Fasta sequences of the BLASTed samples.
        """
        self.json_list = json_list
        self.jsonfolder = jsonfolder
        # Dict for easy search by name as index
        self.fasta_seqs = self.__get_fasta_dict(fasta_data)

        # Script path; MUSCLE.exe and .mao file position
        self.files_path = os.path.realpath('..\\files').replace('\\', '/')

        # Parameters for species name cleanup
        self.species_name_filter = ['subsp','var']
        # BLAST hit filters
        self.min_hits_filter = 5 # Minimum number of hits to keep
        self.max_hits_filter = 20 # Maximum number of hits to keep
        self.min_percentage = 0.94 # %requiered (applied after min_hits_filter)
        # Report table filters
        self.best_hits_span = 0.003 # % of other hits away from best to consider
        self.short_seq_len = 300 # Short sequence to be analysed seperately


    def __get_fasta_dict(self, fasta_data):
        """ Return a dictionnay of samples names and their fasta sequence """
        
        fasta_list = fasta_data.strip('\n').split('\n')
        fasta_seqs = {}
        #for i in range(1, len(fasta_list), 2):
        #   fasta_seqs[fasta_list[i-1].strip('>')] = fasta_list[i-1] + '\n' + fasta_list[i]

        name = ''
        seqs = []
        for i in fasta_list:
            if i:
                if i[0] == '>':
                    if seqs:
                        fasta_seqs[name.strip('>')] = name + '\n' + ''.join(seqs)
                    name = i
                    seqs = []
                else:
                    seqs.append(i)
        fasta_seqs[name.strip('>')] = name + '\n' + ''.join(seqs)

        return fasta_seqs


    def __correct_species_name(self, name):
        """
        Remove any strain name that could be attached to the species name

        Parameters
        ----------
        name : STR
            The unmodified species name from the BLAST results.

        Returns
        -------
        STR
            The Species name with no strain name attached to it.
        """
        splitted_name = name.split()
        if len(splitted_name) <= 2:
            return name

        if splitted_name[2].strip('.') in self.species_name_filter:
            return ' '.join(splitted_name[:4])
        else:
            return ' '.join(splitted_name[:2])


    def __find_best_matches(self, data):
        """
        Finds the best matches according to the BLAST filter parameters
        - unique species name
        - at least self.min_hits_filter results no matter percentage
        - stop when below self.min_percentage
        - stop at self.max_hits_filter results

        Return:
            matches: DICT of all the hits found after appling the filters
                    Key = STR, accession number STR of the hit
                    Value = STR, Fasta seq of the hit
            matches_percentage: List all the hit percentages
                    Value = [Species name, % identity, % coverage]
        """
        species_list = [] # all species seen
        counter = 0
        matches = {}
        matches_percentage = []
        sequence_len = data['report']['results']['search']['query_len']
        for hit in data['report']['results']['search']['hits']:
            # Ignore if species has already been added
            species = self.__correct_species_name(hit['description'][0]['sciname'])
            if species in species_list:
                continue

            identity = hit['hsps'][0]['identity']
            align_len = hit['hsps'][0]['align_len']
            percent_identity = identity/align_len
            cover = align_len/sequence_len
            if cover > 1:
                cover = 1

            # Applying filters
            if (counter >= self.min_hits_filter
                and percent_identity < self.min_percentage) \
                or counter >= self.max_hits_filter:
                break

            # Add species fasta seq in matches
            acc_num = hit['description'][0]['accession']
            if acc_num == 'CP066061': #Ignoring this hit for S.cerevisiae
                continue
            # Added to a dictionnary to prevent adding the same acc_num multiple times
            matches[acc_num] = '>' + species + ' ' + acc_num + '\n' + hit['hsps'][0]['hseq']
            matches_percentage.append([species, percent_identity, cover])
            species_list.append(species)
            counter += 1

        return matches, matches_percentage


    def __parse_all_json(self, json_reports, jsonfolder, fasta_seqs):
        """
        Fetch the BLAST results information for each samples

        Parameters
        ----------
        json_reports : Dict
            contains the BLAST result information for all samples.
        jsonfolder : STR
            Path of the JSON files.
        fasta_seqs : DICT
            key: Name of sequence (str), value: fasta sequence with >name (str)

        Returns
        -------
        results : list of lists
            Each list contains the BLAST info of each samples.
            [name, length, species, identity, accession, fasta_seq, matches]

        """

        # Loops all reports and extract best match from each
        results = []
        for report in json_reports['BlastOutput2']:

            # extract information
            matches, matches_per = self.__find_best_matches(report)
            name = report['report']['results']['search']['query_title']
            length = report['report']['results']['search']['query_len']

            if matches_per:
                best_species = matches_per[0][0]
                best_per_identity =  matches_per[0][1]
                best_cover = matches_per[0][2]
                best_accession = list(matches.keys())[0]
            else:
                best_species = 'No hits found'
                best_per_identity =  0
                best_cover = 0
                best_accession = 'No hits found'

            # save information
            results.append([name,
                            length,
                            best_species,
                            f'{best_per_identity:.2%}',
                            f'{best_cover:.2%}',
                            best_accession,
                            fasta_seqs[name],
                            matches,
                            matches_per])

        return results


    def __newick_tree(self, out_file):
                """
                Uses the aligned fasta file to create a NJ tree
                The tree is saved in newick tree file
                """
                out_tree = f'{out_file[:-4]}_Tree.nwk'
                print(f'NJ tree: {out_tree.split("/")[-1]}')
                mao_specifications = f'{self.files_path}/infer_NJ_nucleotide.mao'
                #NJ command line
                # stream = os.popen(
                #f'megacc -a infer_NJ_nucleotide.mao -d "{out_file}" -o "{out_file[:-4]}.nwk"')
                # output = stream.read()
                # print(output)
                # Add -n to remove summary
                nj_cline = f'megacc -a {mao_specifications} -d "{out_file}" -o "{out_tree}" -n -s'
                os.system(nj_cline)


    def __muscle_align(self, in_file):
            """
            Align a fasta file (in_file) with MUSCLE, output a aligned fasta file
            Then it uses the aligned fasta file to create a NJ tree
            The tree is saved in newick tree file
            """

            #Align sequences with MUSCLE
            muscle_exe = f"{self.files_path}/muscle.exe"
            out_file = f'{in_file[:-4]}_aligned.fas'
            # Creating the command line
            muscle_cline = MuscleCommandline(muscle_exe,
                                             input=in_file,
                                             out=out_file)
            # Lauch the muscle command line
            print(f'MUSCLE alignment: {out_file.split("/")[-1]}')
            muscle_cline()

            # Create a NJ newick  tree
            self.__newick_tree(out_file)

    def parse(self):
        """
        Parse the fasta_data
        """
        # Parsing the BLAST multiple-file JSON results
        self.results = self.__parse_all_json(self.json_list,
                                             self.jsonfolder,
                                             self.fasta_seqs)

        self.__create_dataframe()


    def __create_dataframe(self):
        # Make result list into pandas dataframe
        self.df = pd.DataFrame(self.results, columns=['Sample Name', 'Length',
                                            'Best Species', '% Identity',
                                            '% Cover', 'Acc Num', 'Sequence',
                                            'Matches', 'matches_per'])


    def to_csv(self, path):
        """
        Save the Blast Results into a csv file.

        Parameters
        ----------
        path : STR
            Folder in which the csv file is saved.
        """

        # Adding best matches column, filter according to best_hits_span
        def best_matches_filter(hit):
            if not hit: # No matches
                return []
            max_per = hit[0][1]
            return [i[:2] for i in hit if i[1] >= max_per-self.best_hits_span]

        self.df['best_matches'] = self.df['matches_per'].apply(best_matches_filter)

        # Formating the matches columns for readability
        self.df['Best Matches'] = self.df['best_matches'].apply(
            lambda x:'\n'.join([f'{i[0]} ({i[1]:.2%})' for i in x]))

        self.df['All Matches'] = self.df['matches_per'].apply(
            lambda x:'\n'.join([f'{i[0]} ({i[1]:.2%}) cov:{i[2]:.2%}' for i in x]))

        # Save BLAST results dataframe into a csv file
        columns = ['Sample Name',
                   'Length',
                   'Best Species',
                   '% Identity',
                   '% Cover',
                   'Best Matches',
                   'All Matches']

        if 'Figure' in self.df.columns:
            columns.append('Figure')

        try:
            self.df[columns].to_csv(path, index=False)
        except PermissionError:
            raise Error('BLAST_results.csv file is already open. Please close file.')


    def create_dir_files(self, dirname, do_alignment):
        """
        Create 1 folder per unique genus which contain a fasta file with all
        sample sequences belonging to that genus

        If do_alignment, the closest matches are added to the fasta file and
        an aligned fasta file is saved separatly. Also, a newick tree of the
        alignment is saved.
        """
        
        def prioritized_merging(matches):
            """ 
            Merge the matches of different samples.
            Prioritize the longest sequence when encountering duplicates.
            """
            strains = {}
            for hits in matches:
                for acc_num, fasta_seq in hits.items():
                    if strains.get(acc_num):
                        # Fasta_seq is 2 lines, separated by \n
                        if len(strains[acc_num].split('\n')[1]) < len(fasta_seq):
                            strains[acc_num] = fasta_seq
                    else:
                        strains[acc_num] = fasta_seq
            return strains
        
        
        try:
            # Create a genus field
            self.df['Genus'] = self.df['Best Species'].str.split().apply(lambda x: x[0])
            # Create the main folders
            os.mkdir(f'{dirname}/Sequences/')
            print(f'\nSequences are saved in {dirname}/Genus/')
            if do_alignment:
                os.mkdir(f'{dirname}/Genus/')
                print(f'\nAlignments are saved in {dirname}/Genus/')
            
            for genus in self.df['Genus'].unique():
                print('------------------------------------------------')

                # Save all samples for which no hit were found
                if genus == 'No':
                    seq_all = self.df.loc[(self.df['Genus'] == genus)]['Sequence']
                    # Save sequences in a txt file with the genus name
                    with open(f'{dirname}/Sequences/No_hits_found.fas', 'w') as file:
                        file.write('\n'.join(seq_all.to_list()))
                        print('File created: /Sequences/No_hits_found.fas')
                    continue

                # Make a folder for current genus
                if do_alignment:
                    os.mkdir(f'{dirname}/Genus/{genus}')
                    print(f'Folder created: {genus}')
                # Apply the short sequence filter
                dfall = self.df.loc[(self.df['Genus'] == genus)]
                dflong = self.df.loc[(self.df['Genus'] == genus) &
                                     (self.df['Length'] > self.short_seq_len)]
                dfshort = self.df.loc[(self.df['Genus'] == genus) &
                                      (self.df['Length'] <= self.short_seq_len)]
                seq_all = dfall['Sequence']
                seq_genus = dflong['Sequence']
                seq_genus_short = dfshort['Sequence']
                # Add hits for all samples of same genus in a dictionnary
                # Duplicates will overwrite each other since acc_num is used as unique key
                #type_strains_long = {k:v for d in dflong['Matches'] for k, v in d.items()}
                #type_strains_short = {k:v for d in dfshort['Matches'] for k, v in d.items()}
                type_strains_long = prioritized_merging(dflong['Matches'])
                type_strains_short = prioritized_merging(dfshort['Matches'])
            
                # Save sequences in a txt file with the genus name
                with open(f'{dirname}/Sequences/{genus}.fas', 'w') as file:
                    file.write('\n'.join(seq_all.to_list()))
                    print(f'File created: /Sequences/{genus}.fas')

                # Save sequences with type strains and do alignement + Tree
                if len(seq_genus) > 0 and do_alignment:
                    with open(f'{dirname}/Genus/{genus}/{genus}_wTS.fas', 'w') as file:
                        file.write('\n'.join(seq_genus.to_list()))
                        file.write('\n')
                        file.write('\n'.join(type_strains_long.values()))
                        print(f'File created: Genus/{genus}/{genus}_wTS.fas')
                    # MUSCLE alignment of the sequences, saving in a new fasta file
                    self.__muscle_align(f'{dirname}/Genus/{genus}/{genus}_wTS.fas')

                # Save short sequences with type strains and do alignement + Tree
                if len(seq_genus_short) > 0 and do_alignment:
                    # Save the sequences in fasta format for short sequence
                    with open(f'{dirname}/Genus/{genus}/{genus}_short_wTS.fas', 'w') as file:
                        file.write('\n'.join(seq_genus_short.to_list()))
                        file.write('\n')
                        file.write('\n'.join(type_strains_short.values()))
                        print(f'File created: Genus/{genus}/{genus}_short.fas')
                    # MUSCLE alignment of the sequences, saving in a new fasta file
                    self.__muscle_align(f'{dirname}/Genus/{genus}/{genus}_short_wTS.fas')

        except OSError:
            raise Error("Creation of the directories or files failed." +
                        " The script does not overwrite existing folders.")
        else:
            print('------------------------------------------------')
            print ("Successfully created all directories and files")


    def create_figure_file(self, path):


        def get_para_data(output_doc_name, paragraph):
            """
            For copying a paragraph with all its characteristics
            No true copy command exist, so manually it is!
            All runs of a paragraph need to be copied individually

            Input:
                paragraph: The paragraph to transfer
                output_doc_name : The Document in which the paragraph is to be copied
            """
            # Creating a new paragraph (output) in the targeted Document
            output_para = output_doc_name.add_paragraph()
            # Copying all the runs of the input paragraph to the output paragraph
            for run in paragraph.runs:
                output_run = output_para.add_run(run.text)
                output_run.bold = run.bold
                output_run.italic = run.italic
                output_run.underline = run.underline
                output_run.font.color.rgb = run.font.color.rgb
                output_run.style.name = run.style.name
                output_run.font.size = run.font.size
                output_run.font.superscript = run.font.superscript
            # Copying Paragraph's format data
            output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
            output_para.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
            #output_para.paragraph_format.element = paragraph.paragraph_format.element
            output_para.paragraph_format.keep_together = paragraph.paragraph_format.keep_together
            output_para.paragraph_format.keep_with_next = paragraph.paragraph_format.keep_with_next
            output_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
            output_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
            output_para.paragraph_format.line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
            output_para.paragraph_format.page_break_before = paragraph.paragraph_format.page_break_before
            output_para.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
            output_para.paragraph_format.space_after = paragraph.paragraph_format.space_after
            output_para.paragraph_format.space_before = paragraph.paragraph_format.space_before
            #output_para.paragraph_format.tab_stops = paragraph.paragraph_format.tab_stops
            output_para.paragraph_format.widow_control = paragraph.paragraph_format.widow_control


        def copy_doc(to_doc, from_doc):
            """
            Copy a whole document (copy each paragraph)

            Inputs:
                to_doc: Document in which all paragraph are to be copied
                        New paragraph are added at the end of the already existing ones
                from_doc: Document where all paragraph are to be copied.
            """
            for para in from_doc.paragraphs:
                get_para_data(to_doc, para)


        genus_count = self.df['Genus'].value_counts()
        genera = self.df['Genus'].unique()
        genera.sort()
        figure_counter = 1
        table = []
        for genus in genera:
            if genus == 'No': # No matches
                table.append([genus, genus_count[genus], 0])
                continue
            table.append([genus, genus_count[genus], figure_counter])
            figure_counter +=1

        # Convert figure numbers into a dictionnary for easy access
        genus_figures = {i[0]:i[2] for i in table}

        # Save the genus count into a .csv file
        dfgenus = pd.DataFrame(table, columns=['Genus','Count','Figure'])
        dfgenus.to_csv(path + '/genus_count.csv', index=False)
        print('\nThe count of each genus was saved in "genus_count.csv".')

        # Create the final Document from the template
        final_doc = docx.Document(f'{self.files_path}/template.docx')

        # Create new column with the figure number
        self.df['Figure'] = 0 # Initiating
        self.df['Figure'] = self.df['Genus'].apply(lambda x: genus_figures[x])

        for genus in genus_figures:
            if genus == 'No': # No matches
                continue
            # Create a temp doc containing one figure, created with the template
            temp_doc = docx.Document(f'{self.files_path}/template.docx')
            # Modify
            modif_dict = {'&1': str(genus_figures[genus]),
                          '&2': genus}

            for modif in modif_dict:
                for run in temp_doc.paragraphs[0].runs:
                    if modif in run.text:
                        run.text = run.text.replace(modif, modif_dict[modif])
            # temp_doc.paragraphs[0].runs[0].text = f'Figure {genus_figures[genus]}.'
            # temp_doc.paragraphs[0].runs[5].text = genus + ' '
            # Save
            copy_doc(final_doc, temp_doc)

        # Save the final doc
        final_doc.save(path + '/figures_template.docx')
        print('\nA figure template was created and saved in "figures_template.docx".')


    def display_parameters(self):
        """
        Display the parameters for filtering of results.
        """

        print(f"""\n
---------------------------------------------------------
Current analysis parameters:

    BLAST hit filters:
        Minimum number of hits to keep: {self.min_hits_filter} hits
        Maximum number of hits to keep: {self.max_hits_filter} hits
        then
        Minimum percentage of ID required: {self.min_percentage}

    Report table filters:
        Genetic distance % for Best Matches to be accepted: {self.best_hits_span}
        Short sequences threshold: {self.short_seq_len} bp
            (recommended: 300bp for 26S and 400 bp for 16S)
---------------------------------------------------------
""")


    def change_parameters(self):
        """
        Menu option to modify the parameters for the filtering of results.
        """
        while True:
            print('\nSelect parameter:')
            print(f'1) Minimum number of hits to keep: {self.min_hits_filter} hits')
            print(f'2) Maximum number of hits to keep: {self.max_hits_filter} hits')
            print(f'3) Minimum percentage required: {self.min_percentage}')
            print(f'4) Genetic distance % for Best Matches: {self.best_hits_span}')
            print(f'5) Short sequences threshold: {self.short_seq_len} bp')
            print('6) Back\n')
            query = input('Choice: ')

            if query == '6':
                break

            new_value = input('Enter the new parameter value: ')

            try:
                new_value = float(new_value)
            except ValueError:
                print('Not a number!')
                continue

            if query == '1':
                if new_value < 1 or new_value >= self.max_hits_filter:
                    print('Invalid value. No change.')
                    print(f'Valid range: 1 to {self.max_hits_filter - 1}')
                else:
                    self.min_hits_filter = int(new_value)
            elif query == '2':
                if new_value <= self.min_hits_filter or new_value >= 1000:
                    print('Invalid value. No change.')
                    print(f'Valid range: {self.min_hits_filter + 1} to 1000')
                else:
                    self.max_hits_filter = int(new_value)
            elif query == '3':
                if new_value < 0 or new_value > 1:
                    print('Invalid value. No change.')
                    print('Valid range: 0 to 1.0 ')
                else:
                    self.min_percentage = new_value
            elif query == '4':
                if new_value < 0 or new_value > 0.1:
                    print('Invalid value. No change.')
                    print('Valid range: 0 to 0.1')
                else:
                    self.best_hits_span = new_value
            elif query == '5':
                if new_value <= 0:
                    print('Invalid value. No change.')
                    print('Valid range: minimun of 1')
                else:
                    self.short_seq_len = new_value
            elif not query:
                continue


class Error(Exception):
    """Exception raised for errors in the input.

    Attributes:
        message -- explanation of the error
    """

    def __init__(self,message):
        self.message = 'ERROR: '+message

    def __str__(self):
        return self.message



# ----------------------------------------------------------------------------


print('*Do not execute this script on files that are saved on the P:/G: '+
      'server since it creates new files and folders.*\n\n')
print("""INSTRUCTION:
Place all 16S or 26/28S sequences in a single file in FASTA format then
BLAST all sequences at the same time by loading in the FASTA file on the
NCBI BLAST Website.

For 16S, select the 16S database from the rRNA/ITS Databases section.
No need to check the 'Sequence from type material' button.

For 26S/28S, select the nucleotide collection from the Standard Database
and check the 'Sequence from type material' button.

Once the BLAST search is done, download the 'Single-file JSON' from the
download dropdown menu.\n\n""")
input('Press Enter to continue.\n')
print('Select the downloaded Single-File JSON from the BLAST result:')


# Ask for the main JSON file path
json_file_path = filedialog.askopenfilename(
    title="Select an JSON File",
    filetypes=(("JSON files", "*.json"), ("All files", "*.*"))
    )


print('Select the FASTA file used for the BLAST search:')
# Ask for the fasta file path
fasta_file_path = filedialog.askopenfilename(
    title="Select an FASTA main File",
    filetypes=(("FASTA files", "*.fas"), ("All files", "*.*"))
    )

try:
    # Loading main JSON file
    with open(json_file_path) as file:
        json_file = json.load(file)

    # Loading FASTA sequences
    with open(fasta_file_path) as fasta_file:
        fasta_data = fasta_file.read()

except FileNotFoundError:
    print('ERROR: File selection was cancelled.')
    input("Press Enter to exit.")
    exit()

try:
    json_folder = ntpath.dirname(json_file_path) # Folder of the json file
    fasta_folder = ntpath.dirname(fasta_file_path) # Folder of the fasta file

    # Initializing
    blastdata = BlastParse(json_file, json_folder, fasta_data)

    blastdata.display_parameters()

    while True:
        print('Required analysis:')
        print('1) Save the BLAST result in a table.')
        print('2) Previous + Fasta sequences in genus specific files.')
        print('3) Previous + alignement with closest TYPE strains and NJ tree.')
        print('4) Previous + figure headers for report with added figure link.')
        print('5) Change parameters.')
        query = input('Choice: ')

        if query.lower() == '5':
            blastdata.change_parameters()
        else:
            break

    # Parsing the Blast data
    blastdata.parse()

    if query.lower() == '2':
        blastdata.create_dir_files(fasta_folder, do_alignment=False) # Has print()
    elif query.lower() == '3':
        blastdata.create_dir_files(fasta_folder, do_alignment=True) # Has print()
    elif query.lower() == '4':
        blastdata.create_dir_files(fasta_folder, do_alignment=True)
        blastdata.create_figure_file(fasta_folder)

    # Save in .csv file
    blastdata.to_csv(f'{fasta_folder}/BLAST_results.csv')
    print(f'\nResults summary saved in {fasta_folder}/BLAST_results.csv')
    print()

except Error as err:
    print(err)
finally:
    input("Press Enter to exit.")
