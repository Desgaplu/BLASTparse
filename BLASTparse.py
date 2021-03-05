#! python3
# -*- coding: utf-8

"""
BLASTparse.py
Created on Mon Dec 14 14:54:15 2020

Descrition:
Parse a multiple JSON BLAST output for the best match of each sequence and
uses the fasta sequences of the samples to classify them in folder names with
the closest genus identified with the BLAST search.

Output a TSV file containing the sample name, samples seq length,
closest species, %identity, species accession number and the sample fasta
sequence for each samples.
It then filter the TSV file to create a folder for each genus and add a fasta
file containing all sample sequences belonging to that genus.

@author: pldesgagne
"""

import tkinter as tk  # using the open filedialog
from tkinter import filedialog  # using the open filedialog
import ntpath  # get the filename from path string
import json
import os
import pandas as pd
from sys import exit
from Bio.Align.Applications import MuscleCommandline
import docx

# Muscle.exe need to be in the same folder as this script
# from http://www.drive5.com/muscle/downloads.htm

# MEGACC (MEGA command console) need to be installed on computer
# The script use megacc command line from os.system() to create NJ trees
# The parameters of the tree creation are saved in file infer_NJ_nucleotide.mao
# This file is created with MEGAX in prototype mode


"""
TODO:
    -Ajust parameters for the selection of closest match
        -Current:   -Unique species name
                    -At least 5 hits
                    -Stop when %identity < 95%
                    -Max of 20 hits

                    -cutoff at 300bp
                    -best results range: 0.3%-0.5% away from best result
                        -If more than 5, write "or other closely related species"

    -Ajust length (current:300bp) for what is considered a short sequence

    -Add an option to change parameters

    -Add a function that create a word document with genus figure names

    -Create a csv file with the species count

    -Create better instructions with a pause before the files query

    -Change the 'Species' column to 'Best Species'

    -Create 2 output tables, one for final report table and another for
        all raw results.
        For final table, see TotalTableParse.py.
            Maybe add option to create .doc with figure and add figure num
            to this table.

    -Add other best results within 0.3% of best percentage

    -Pick best results from the results table returned from function
        instead of doing 2 calculations

    -Change the script into an easily readable format
        -Convert the script to execute all functions at same level
        for better readability
        -If requiered, make functions fully inclosed other than global vars

    -Add the species name correction from the MEGA tree name fixer script
        when unique species name are added to the species_list in
        find_best_matches function

    -Ktinker GUI?

DONE:
    -Add output options:
        -Summary only
        -Fasta sequences separated by genus
        -Fasta sequences separated by genus aligned with closest TYPE strain +
            newick NJ tree output.
        -Add the species rename function similar to MEGAparse.py script

    -Add percentage of identity on the 5 closest matches
        -All percentages were added


"""

# Initialise tkinter to enable the uses of filedialog
root = tk.Tk()
root.withdraw()  # Prevent a empty window to be opened


class BlastParse():
    """
    Parse the multiple-json output of BLASTn with multiple 16S or 26S
    rDNA sequences and arrange them in a table.

    It is also possible to align all the sequences of the same genus
    with their best hits and create a NJ tree with the output. Those are
    saved in individual folder (named with the genus).

    Parameters
        ----------
        json_list : DICT
            Contain a list of the name of each json file.
        jsonfolder : STR
            Path to the json files.
        fasta_seqs : DICT
            Fasta sequences of the BLASTed samples.
            Key= Sample name STR
            Value= Fasta seq STR
        species_name_filter : LIST of STR
            keyword for the clean up of species names
        min_hits_filter : STR
            Minimum number of possible hits
        max_hits_filter : STR
            Maximum number of possible hits
        min_percentage : STR
            Minimum id percentage accepted for a hit
        best_hits_span : FLOAT
            Id percentage distance away from best hit for acceptable top hits
        short_seq_len : INT
            Len of the sequence under which the alignement is done separately.
        results : List of list
            Compiled results from all JSON files.
    """

    def __init__(self, json_list, jsonfolder, fasta_data):
        """
        Create the parser with default parameters.

        Parameters
        ----------
        json_list : DICT
            Contain a list of the name of each json file.
        jsonfolder : STR
            Path to the json files.
        fasta_seqs : STR
            Fasta sequences of the BLASTed samples.
        """
        self.json_list = json_list
        self.jsonfolder = jsonfolder
        # Dict for easy search by name as index
        self.fasta_seqs = self.__get_fasta_dict(fasta_data)

        # Script path; MUSCLE.exe and .mao file position
        self.dirname = os.getcwd().replace('\\', '/')
        # Parameters for species name cleanup
        self.species_name_filter = ['subsp','var']
        # BLAST hit filters
        self.min_hits_filter = 5 # Minimum number of hits to keep
        self.max_hits_filter = 20 # Maximum number of hits to keep
        self.min_percentage = 0.94 # %requiered (applied after min_hits_filter)
        # Report table filters
        self.best_hits_span = 0.003 # % of other hits away from best to consider
        self.short_seq_len = 300 # Short sequence to be analysed seperately


    def __get_fasta_dict(self, fasta_data):
        """ Return a dictionnay of samples names and their fasta sequence """
        fasta_list = fasta_data.strip('\n').split('\n')
        fasta_seqs = {}
        for i in range(1, len(fasta_list), 2):
            fasta_seqs[fasta_list[i-1].strip('>')] = fasta_list[i-1] + '\n' + fasta_list[i]

        return fasta_seqs


    def __correct_species_name(self, name):
        """
        Remove any strain name that could be attached to the species name

        Parameters
        ----------
        name : STR
            The unmodified species name from the BLAST results.

        Returns
        -------
        STR
            The Species name with no strain name attached to it.
        """
        splitted_name = name.split()
        if len(splitted_name) <= 2:
            return name

        if splitted_name[2].strip('.') in self.species_name_filter:
            return ' '.join(splitted_name[:4])
        else:
            return ' '.join(splitted_name[:2])


    def __find_best_matches(self, data):
        """
        Finds the best matches according to the BLAST filter parameters
        - unique species name
        - at least self.min_hits_filter results no matter percentage
        - stop when below self.min_percentage
        - stop at self.max_hits_filter results

        Return:
            matches: DICT of all the hits found after appling the filters
                    Key = accession number STR of the hit
                    Value = Fasta seq of the hit
            matches_percentage: List all the hit percentages
                    Value = [Species name, % identity, % coverage]
        """
        species_list = [] # all species seen
        counter = 0
        matches = {}
        matches_percentage = []
        for hit in data['BlastOutput2']['report']['results']['search']['hits']:
            # Ignore if species has already been added
            species = self.__correct_species_name(hit['description'][0]['sciname'])
            if species in species_list:
                continue

            query_to = hit['hsps'][0]['query_to']
            identity = hit['hsps'][0]['identity']
            align_len = hit['hsps'][0]['align_len']
            per_identity = identity/align_len
            cover = align_len/query_to
            if cover > 1:
                cover = 1

            # Applying filters
            if (counter >= self.min_hits_filter
                and per_identity < self.min_percentage) \
                or counter >= self.max_hits_filter:
                break

            # Add species fasta seq in matches
            acc_num = hit['description'][0]['accession']
            # Added to a dictionnary to prevent adding the same acc_num multiple times
            matches[acc_num] = '>' + species + ' ' + acc_num + '\n' + hit['hsps'][0]['hseq']
            matches_percentage.append([species, per_identity, cover])
            species_list.append(species)
            counter += 1

        return matches, matches_percentage


    def __parse_all_json(self, json_list, jsonfolder, fasta_seqs):
        """
        Fetch the BLAST results information for each samples (1 per JSON file)

        Parameters
        ----------
        data_list : JSON
            contains a list of JSON file names.
        jsonfolder : STR
            Path of the JSON files.
        fasta_seqs : DICT
            key: Name of sequence (str), value: fasta sequence with >name (str)

        Returns
        -------
        results : list of lists
            Each list contains the BLAST info of each samples.
            [name, length, species, identity, accession, fasta_seq, matches]

        """

        # Loops all file and extract best match from each
        results = []
        for file in json_list['BlastJSON']:
            # load file
            with open(f"{jsonfolder}/{file['File']}") as json_file:
                data = json.load(json_file)
            # extract information
            matches, matches_per = self.__find_best_matches(data)
            name = data['BlastOutput2']['report']['results']['search']['query_title']
            length = data['BlastOutput2']['report']['results']['search']['query_len']

            best_species = matches_per[0][0]
            best_per_identity =  matches_per[0][1]
            best_cover = matches_per[0][2]
            best_accession = list(matches.keys())[0]


            # save information
            results.append([name,
                            length,
                            best_species,
                            f'{best_per_identity:.2%}',
                            f'{best_cover:.2%}',
                            best_accession,
                            fasta_seqs[name],
                            matches,
                            matches_per])

        return results


    def __newick_tree(self, out_file):
                """
                Uses the aligned fasta file to create a NJ tree
                The tree is saved in newick tree file
                """
                out_tree = f'{out_file[:-4]}_Tree.nwk'
                print(f'NJ tree: {out_tree.split("/")[-1]}')
                mao_specifications = f'{self.dirname}/infer_NJ_nucleotide.mao'
                #NJ command line
                # stream = os.popen(
                #f'megacc -a infer_NJ_nucleotide.mao -d "{out_file}" -o "{out_file[:-4]}.nwk"')
                # output = stream.read()
                # print(output)
                # Add -n to remove summary
                nj_cline = f'megacc -a {mao_specifications} -d "{out_file}" -o "{out_tree}" -n -s'
                os.system(nj_cline)


    def __muscle_align(self, in_file):
            """
            Align a fasta file (in_file) with MUSCLE, output a aligned fasta file
            Then it uses the aligned fasta file to create a NJ tree
            The tree is saved in newick tree file
            """

            #Align sequences with MUSCLE
            muscle_exe = f"{self.dirname}/muscle.exe"
            out_file = f'{in_file[:-4]}_aligned.fas'
            # Creating the command line
            muscle_cline = MuscleCommandline(muscle_exe,
                                             input=in_file,
                                             out=out_file)
            # Lauch the muscle command line
            print(f'MUSCLE alignment: {out_file.split("/")[-1]}')
            muscle_cline()

            # Create a NJ newick  tree
            self.__newick_tree(out_file)

    def parse(self):
        """
        Parse the fasta_data
        """
        # Parsing the BLAST multiple-file JSON results
        self.results = self.__parse_all_json(self.json_list,
                                             self.jsonfolder,
                                             self.fasta_seqs)

    def to_csv(self, path):
        """
        Save the Blast Results into a csv file.

        Parameters
        ----------
        path : STR
            Folder in which the csv file is saved.
        """

        # Make into dataframe (to use Dataframe.to_csv function)
        self.df = pd.DataFrame(self.results, columns=['Name', 'Length',
                                            'Best Species', '% Identity',
                                            '% Cover', 'Acc Num', 'Sequence',
                                            'Matches', 'matches_per'])
        # Adding best matches column
        def best_matches_filter(hit):
            max_per = hit[0][1]
            return [i[:2] for i in hit if i[1] >= max_per-self.best_hits_span]

        self.df['best_matches'] = self.df['matches_per'].apply(best_matches_filter)

        # Formating the matches columns for readability
        self.df['Best Matches'] = self.df['best_matches'].apply(
            lambda x:'\n'.join([f'{i[0]} ({i[1]:.2%})' for i in x]))

        self.df['All Matches'] = self.df['matches_per'].apply(
            lambda x:'\n'.join([f'{i[0]} ({i[1]:.2%}) cov:{i[2]:.2%}' for i in x]))

        # Save BLAST results dataframe into a csv file
        columns = ['Name',
                   'Length',
                   'Best Species',
                   '% Identity',
                   '% Cover',
                   'Best Matches',
                   'All Matches']
        try:
            self.df[columns].to_csv(path, index=False)
        except PermissionError:
            raise Error('BLAST_results.csv file is already open. Please close file.')


    def create_dir_files(self, dirname, do_alignment):
        """
        Create 1 folder per unique genus which contain a fasta file with all
        sample sequences belonging to that genus

        If do_alignment, the closest matches are added to the fasta file and
        an aligned fasta file is saved separatly. Also, a newick tree of the
        alignment is saved.
        """

        try:
            # Create a genus field
            self.df['Genus'] = self.df['Best Species'].str.split().apply(lambda x: x[0])
            # Create the main folders
            os.mkdir(f'{dirname}/Sequences/')
            print(f'\nSequences are saved in {dirname}/Genus/')
            if do_alignment:
                os.mkdir(f'{dirname}/Genus/')
                print(f'\nAlignments are saved in {dirname}/Genus/')
            # For each genus
            for genus in self.df['Genus'].unique():
                print('------------------------------------------------')
                # Make a folder for that genus
                if do_alignment:
                    os.mkdir(f'{dirname}/Genus/{genus}')
                    print(f'Folder created: {genus}')
                # Filter all sequences of that genus
                dfall = self.df.loc[(self.df['Genus'] == genus)]
                dflong = self.df.loc[(self.df['Genus'] == genus) &
                                     (self.df['Length'] > self.short_seq_len)]
                dfshort = self.df.loc[(self.df['Genus'] == genus) &
                                      (self.df['Length'] <= self.short_seq_len)]
                seq_all = dfall['Sequence']
                seq_genus = dflong['Sequence']
                seq_genus_short = dfshort['Sequence']
                # Add matches for all samples of same genus in a dictionnary
                # Duplicates will overwrite each other since acc_num is used as unique key
                type_strains_long = {k:v for d in dflong['Matches'] for k, v in d.items()}
                type_strains_short = {k:v for d in dfshort['Matches'] for k, v in d.items()}
                # =============================================================
                # Possibility to change the attribution of type_strains_long and
                # type_strains_short into an actual fonction that checks the
                # len of the sequence of the possible duplicate already present
                # in the dict to keep the longest sequence. For now, the
                # NJ tree is done with "complete deletion" that create a tree
                # based on the aligned region of all of its sequence which is
                # limited to its shortest sequence. The shortest sequence will
                # most likely be one of the samples. If another type of tree
                # alignement is used, the longest type strain seq will need to be
                # selected (they are already croped by BLAST to fit the
                # BLASTed sequence). Other type of alignement would allow the
                # alignement of long and short sequence together and the resulting
                # tree would better show the differences of each sequences
                # individually to every type strains even if the differences are
                # outside of the general aligned region.
                # =============================================================


                # Save sequences in a txt file with the genus name
                with open(f'{dirname}/Sequences/{genus}.fas', 'w') as file:
                    file.write('\n'.join(seq_all.to_list()))
                    print(f'File created: /Sequences/{genus}.fas')
                    
                # Save sequences with type strains and do alignement + Tree
                if len(seq_genus) > 0 and do_alignment:
                    with open(f'{dirname}/Genus/{genus}/{genus}_wTS.fas', 'w') as file:
                        file.write('\n'.join(seq_genus.to_list()))
                        file.write('\n')
                        file.write('\n'.join(type_strains_long.values()))
                        print(f'File created: Genus/{genus}/{genus}_wTS.fas')
                    # MUSCLE alignment of the sequences, saving in a new fasta file
                    self.__muscle_align(f'{dirname}/Genus/{genus}/{genus}_wTS.fas')
                    
                # Save short sequences with type strains and do alignement + Tree
                if len(seq_genus_short) > 0 and do_alignment:
                    # Save the sequences in fasta format for short sequence
                    with open(f'{dirname}/Genus/{genus}/{genus}_short_wTS.fas', 'w') as file:
                        file.write('\n'.join(seq_genus_short.to_list()))
                        file.write('\n')
                        file.write('\n'.join(type_strains_short.values()))
                        print(f'File created: Genus/{genus}/{genus}_short.fas')
                    # MUSCLE alignment of the sequences, saving in a new fasta file
                    self.__muscle_align(f'{dirname}/Genus/{genus}/{genus}_short_wTS.fas')

        except OSError:
            raise Error("Creation of the directories or files failed." +
                        " The script does not overwrite existing folders.")
        else:
            print('------------------------------------------------')
            print ("Successfully created all directories and files")
            
            
    def create_figure_file(self, path):

        
        def get_para_data(output_doc_name, paragraph):
            """
            For copying a paragraph with all its characteristics
            No true copy command exist, so manually it is!
            All runs of a paragraph need to be copied individually
            
            Input:
                paragraph: The paragraph to transfer
                output_doc_name : The Document in which the paragraph is to be copied
            """
            # Creating a new paragraph (output) in the targeted Document
            output_para = output_doc_name.add_paragraph()
            # Copying all the runs of the input paragraph to the output paragraph
            for run in paragraph.runs:
                output_run = output_para.add_run(run.text)
                output_run.bold = run.bold
                output_run.italic = run.italic
                output_run.underline = run.underline
                output_run.font.color.rgb = run.font.color.rgb
                output_run.style.name = run.style.name
                output_run.font.size = run.font.size
                output_run.font.superscript = run.font.superscript
            # Copying Paragraph's format data
            output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
            output_para.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
            #output_para.paragraph_format.element = paragraph.paragraph_format.element
            output_para.paragraph_format.keep_together = paragraph.paragraph_format.keep_together
            output_para.paragraph_format.keep_with_next = paragraph.paragraph_format.keep_with_next
            output_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
            output_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
            output_para.paragraph_format.line_spacing_rule = paragraph.paragraph_format.line_spacing_rule
            output_para.paragraph_format.page_break_before = paragraph.paragraph_format.page_break_before
            output_para.paragraph_format.right_indent = paragraph.paragraph_format.right_indent
            output_para.paragraph_format.space_after = paragraph.paragraph_format.space_after
            output_para.paragraph_format.space_before = paragraph.paragraph_format.space_before
            #output_para.paragraph_format.tab_stops = paragraph.paragraph_format.tab_stops
            output_para.paragraph_format.widow_control = paragraph.paragraph_format.widow_control
        
        
        def copy_doc(to_doc, from_doc):
            """
            Copy a whole document (copy each paragraph)
            
            Inputs:
                to_doc: Document in which all paragraph are to be copied
                        New paragraph are added at the end of the already existing ones
                from_doc: Document where all paragraph are to be copied.
            """
            for para in from_doc.paragraphs:
                get_para_data(to_doc, para)


        genus_count = self.df['Genus'].value_count()
        genera = self.df['Genus'].unique()
        figure_counter = 1
        table = []
        for genus in genera:
            table.append([genus, genus_count[genus], figure_counter])
            figure_counter +=1
        dfgenus = pd.DataFrame(table, columns=['Genus','Count','Figure'])
        dfgenus.to_csv(path)

        # Create the final Document from the template
        final_doc = docx.Document('files\template.docx')
        
        # Initiating new column with the figure number
        self.df['Figure'] = 0
        
        for genus in table:
            # Create a temp doc containing one figure, created with the template
            temp_doc = docx.Document('template.docx')
            # Modify
            temp_doc.paragraphs[0].runs[0].text = f'Figure {genus[2]}.'
            temp_doc.paragraphs[0].runs[4].text = genus
            # Save
            copy_doc(final_doc, temp_doc)
        
        # TODO
        
        # self.df['Figure'].apply(lambda x: if x ==)
        
        # # Save the final doc
        # final_doc.save('output.docx')
        

        
        # # Save genus list with assosiated figure number
        # df.to_csv('tree_count_figure.csv')
        
        # # Load the sample table
        # dfsamples = pd.read_csv('TotalTableCSV_MOD.csv')
        
        # # Create a new figure column, initiate with empty string.
        # dfsamples['Figure'] = ''
        
        # # Copy the appropriate figure number to each samples
        # for index in dfsamples.index:
        #     dfsamples.loc[index]['Figure'] = df.loc[dfsamples.loc[index]['Genus']]['figure']
        
        # # Save the table with the added figure column
        # dfsamples.to_csv('TotalTableCSV_MOD_MOD.csv')

        
class Error(Exception):
    """Exception raised for errors in the input.

    Attributes:
        message -- explanation of the error
    """

    def __init__(self,message):
        self.message = 'ERROR: '+message

    def __str__(self):
        return self.message



# ----------------------------------------------------------------------------


print('*Do not execute this script on files saved on the P:/G: server since it' +
      ' creates folders.*')
print('Locally unzip the multiple-files JSON BLAST result and fasta sequences.')
print("Select the main JSON file; It's the file with no number at the end.\n")

# Ask for the main JSON file path
json_file_path = filedialog.askopenfilename(
    title="Select an JSON main File",
    filetypes=(("JSON files", "*.json"), ("All files", "*.*"))
    )


print('Select the FASTA file used for the BLAST search:')
# Ask for the fasta file path
fasta_file_path = filedialog.askopenfilename(
    title="Select an FASTA main File",
    filetypes=(("FASTA files", "*.fas"), ("All files", "*.*"))
    )

try:
    # Loading main JSON file
    with open(json_file_path) as json_file:
        json_list = json.load(json_file)

    # Loading FASTA sequences
    with open(fasta_file_path) as fasta_file:
        fasta_data = fasta_file.read()

except FileNotFoundError:
    print('ERROR: File selection was cancelled.')
    input("Press Enter to exit.")
    exit()

try:
    json_folder = ntpath.dirname(json_file_path) # Folder of the json file
    fasta_folder = ntpath.dirname(fasta_file_path) # Folder of the fasta file

    # Parsing the Blast data
    blastdata = BlastParse(json_list, json_folder, fasta_data)
    blastdata.parse()
    # Save in .csv file
    blastdata.to_csv(f'{fasta_folder}/BLAST_results.csv')

    print(f'\n\tResults summary saved in {fasta_folder}/BLAST_results.csv')
    print()
    print('Do you wish to classify the sequences in genus specific files/folders?')
    print('1) Fasta sequences in genus specific files')
    print('2) Previous + alignement with closest TYPE strains and NJ tree.')
    print('3) Previous + figure headers for report with added figure link.(TODO)')
    print('4) Change parameters (TODO)')
    print('x) Skip and Exit')
    query = input('Choice: ')
    if query.lower() == '1':
        # Create folders and files with fasta sequences, alignments and TYPE strain acc nums
        blastdata.create_dir_files(fasta_folder, do_alignment=False) # Has print()
    elif query.lower() == '2':
        blastdata.create_dir_files(fasta_folder, do_alignment=True) # Has print()

except Error as err:
    print(err)
finally:
    input("Press Enter to exit.")
